from __future__ import annotations

import csv
import io
import re
from datetime import datetime
from typing import Any


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}
SUPPORTED_DOCUMENT_EXTENSIONS = {".txt", ".docx", ".pdf"} | IMAGE_EXTENSIONS
SUPPORTED_TRANSACTION_EXTENSIONS = {".csv", ".xlsx", ".xlsm", ".pdf"} | IMAGE_EXTENSIONS


def _decode_text(data: bytes) -> str:
    return data.decode("utf-8-sig")


def extract_document_text(data: bytes, *, filename: str) -> str:
    """Extract searchable text from TXT, DOCX, text-based PDF, scanned PDF, or image bytes."""
    suffix = _suffix(filename)
    if suffix == ".txt":
        text = _decode_text(data)
    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - environment-specific
            raise ValueError("DOCX识别需要安装 python-docx") from exc
        document = Document(io.BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                text += "\n" + "\t".join(cell.text.strip() for cell in row.cells)
    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - environment-specific
            raise ValueError("PDF识别需要安装 pypdf") from exc
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if not text.replace("\x00", "").strip():
            # Scanned PDF without text layer: automatically fallback to local OCR
            from legal_funds_agent.parsers.ocr_service import extract_text_from_scanned_pdf
            text = extract_text_from_scanned_pdf(data)
    elif suffix in IMAGE_EXTENSIONS:
        from legal_funds_agent.parsers.ocr_service import extract_text_from_image
        text = extract_text_from_image(data)
    else:
        raise ValueError(f"不支持的文书格式: {suffix or '无扩展名'}")
    text = text.replace("\x00", "").strip()
    if not text:
        raise ValueError("未能从文件中识别出有效文字内容；扫描件或图片请确保清晰且包含可识别文字（OCR）")
    return text


def extract_transactions_csv(data: bytes, *, filename: str) -> str:
    """Return canonical transaction CSV from CSV, common bank XLSX layouts, or transfer screenshots."""
    suffix = _suffix(filename)
    if suffix == ".csv":
        text = _decode_text(data)
        if not text.strip():
            raise ValueError("CSV文件为空")
        return text
    if suffix in IMAGE_EXTENSIONS:
        from legal_funds_agent.parsers.ocr_service import extract_text_from_image, parse_screenshot_transaction
        ocr_text = extract_text_from_image(data)
        parsed = parse_screenshot_transaction(ocr_text)
        if not parsed:
            raise ValueError("未能从转账截图中识别出有效交易要素（金额/收款人）")
        time_val = parsed.get("time", "")
        date_str = time_val[:10] if len(time_val) >= 10 else "2026-03-15"
        time_str = time_val[11:19] if len(time_val) >= 19 else "12:00:00"
        return (
            "transaction_id,date,time,payer,payer_account,payee,payee_account,amount,remark\n"
            f"{parsed['transaction_id']},{date_str},{time_str},被害人,,{parsed['payee']},,{parsed['amount']},转账截图识别\n"
        )
    if suffix == ".pdf":
        return extract_bank_pdf_transactions(data)
    if suffix not in {".xlsx", ".xlsm"}:
        raise ValueError(f"不支持的流水格式: {suffix or '无扩展名'}")
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover - environment-specific
        raise ValueError("XLSX识别需要安装 openpyxl") from exc
    workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheet_records: list[tuple[Any, list[tuple[Any, ...]], dict[str, str], int, dict[str, int]]] = []
    account_registry: dict[str, dict[str, str]] = {}
    for worksheet in workbook.worksheets:
        values = list(worksheet.iter_rows(values_only=True))
        metadata = _sheet_metadata(values)
        header_index = _find_header(values)
        if header_index is None:
            continue
        headers = [_clean_cell(value) for value in values[header_index]]
        mapping = _column_mapping(headers)
        required = {"date_time", "serial", "direction", "counterparty", "amount"}
        if not required.issubset(mapping):
            continue
        sheet_records.append((worksheet, values, metadata, header_index, mapping))
        account_id = metadata.get("account_id", "")
        account_number = metadata.get("account_number", "")
        if account_id:
            account_registry[account_id] = {
                "name": metadata.get("owner_name") or worksheet.title,
                "account_number": account_number,
            }

    rows: list[dict[str, Any]] = []
    for worksheet, values, metadata, header_index, mapping in sheet_records:
        headers = [_clean_cell(value) for value in values[header_index]]
        mapping = _column_mapping(headers)
        required = {"date_time", "serial", "direction", "counterparty", "amount"}
        if not required.issubset(mapping):
            continue
        owner_name = metadata.get("owner_name") or worksheet.title
        owner_account = metadata.get("account_id") or metadata.get("account_number") or ""
        for row_number, raw in enumerate(values[header_index + 1 :], start=header_index + 2):
            if not any(value not in (None, "") for value in raw):
                continue
            direction = _clean_cell(raw[mapping["direction"]])
            if direction not in {"收入", "支出", "入", "出", "贷", "借"}:
                continue
            parsed_time = _parse_datetime(raw[mapping["date_time"]])
            if parsed_time is None:
                continue
            counterparty = _clean_cell(raw[mapping["counterparty"]])
            amount = _number(raw[mapping["amount"]])
            if not counterparty or amount is None:
                continue
            incoming = direction in {"收入", "入", "贷"}
            counterparty_info = _parse_counterparty(counterparty, account_registry)
            counterparty_name = counterparty_info["name"] or counterparty
            counterparty_account = counterparty_info["account"]
            counterparty_account_id = counterparty_info["account_id"]
            payee = owner_name if incoming else counterparty_name
            payer = counterparty_name if incoming else owner_name
            rows.append({
                "transaction_id": _clean_cell(raw[mapping["serial"]]),
                "date": parsed_time.strftime("%Y-%m-%d"),
                "time": parsed_time.strftime("%H:%M:%S"),
                "payer": payer,
                "payer_account": counterparty_account if incoming else owner_account,
                "payer_account_id": counterparty_account_id if incoming else metadata.get("account_id", ""),
                "payee": payee,
                "payee_account": owner_account if incoming else counterparty_account,
                "payee_account_id": metadata.get("account_id", "") if incoming else counterparty_account_id,
                "amount": f"{amount:.2f}",
                "remark": _clean_cell(raw[mapping["remark"]]) if "remark" in mapping else "",
                "source_account_id": metadata.get("account_id", ""),
                # Preserve the original workbook row for later source-file review.
                "source_row": str(row_number),
            })
    if not rows:
        raise ValueError("XLSX中未识别到符合规范的银行流水工作表")
    output = io.StringIO(newline="")
    writer = csv.DictWriter(output, fieldnames=["transaction_id", "date", "time", "payer", "payer_account", "payee", "payee_account", "amount", "remark", "payer_account_id", "payee_account_id", "source_account_id", "source_row"])
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue()


def _suffix(filename: str) -> str:
    match = re.search(r"\.[^.]+$", filename.lower())
    return match.group(0) if match else ""


def _clean_cell(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _find_header(values: list[tuple[Any, ...]]) -> int | None:
    for index, row in enumerate(values):
        normalized = {_clean_cell(value) for value in row}
        if {"交易时间", "银行流水号", "收/支"}.issubset(normalized):
            return index
    return None


def _column_mapping(headers: list[str]) -> dict[str, int]:
    aliases = {
        "date_time": {"交易时间", "交易日期", "日期", "记账日期", "记账时间"},
        "serial": {"银行流水号", "流水号", "交易号", "交易流水号", "凭证号", "业务单号"},
        "direction": {"收/支", "收支", "借贷方向", "借贷", "借/贷", "方向"},
        "counterparty": {"对方户名/账户", "对方账户", "交易对手", "对方名称", "对方户名", "对方姓名"},
        "amount": {"金额", "交易金额", "发生额", "交易金额(元)", "记账金额"},
        "remark": {"摘要/备注", "摘要", "备注", "用途", "交易附言", "交易摘要"},
    }
    return {key: next(index for index, header in enumerate(headers) if header in names) for key, names in aliases.items() if any(header in names for header in headers)}


def _sheet_metadata(values: list[tuple[Any, ...]]) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for row in values[:4]:
        cells = [_clean_cell(value) for value in row]
        for index, cell in enumerate(cells):
            if cell in {"户名", "账户名称"} and index + 1 < len(cells):
                metadata["owner_name"] = cells[index + 1]
            if cell in {"账户ID", "账户编号"} and index + 1 < len(cells):
                metadata["account_id"] = cells[index + 1]
            if cell in {"账号", "账号号码"} and index + 1 < len(cells):
                metadata["account_number"] = cells[index + 1]
    return metadata


def _parse_datetime(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return value
    text = _clean_cell(value).replace("/", "-")
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    return None


def _number(value: Any) -> float | None:
    try:
        return float(str(value).replace(",", "").strip())
    except (TypeError, ValueError):
        return None


def _account_id(value: str) -> str:
    match = re.search(r"账户ID\s*([A-Za-z0-9_-]+)", value)
    return match.group(1) if match else ""


def _parse_counterparty(value: str, account_registry: dict[str, dict[str, str]]) -> dict[str, str]:
    """Split a bank's `name (bank account)` display into canonical fields."""
    text = _clean_cell(value)
    name_match = re.match(r"^(.*?)\s*[（(].*[）)]\s*$", text)
    name = (name_match.group(1) if name_match else text).strip()
    account_text = ""
    if name_match:
        account_text = text[name_match.end(1):].strip()
        account_text = re.sub(r"^[（(]\s*|[）)]\s*$", "", account_text).strip()
    explicit_id = _account_id(text)
    matched_id = explicit_id
    matched_account = ""
    normalized_counterparty = re.sub(r"\s+", "", account_text).upper()
    if not matched_id:
        for account_id, entry in account_registry.items():
            normalized_account = re.sub(r"\s+", "", entry.get("account_number", "")).upper()
            if normalized_account and normalized_account in normalized_counterparty:
                matched_id = account_id
                matched_account = entry.get("account_number", "")
                break
    if not matched_account and account_text:
        matched_account = _extract_account_token(account_text)
    return {"name": name, "account": matched_account, "account_id": matched_id}


def _extract_account_token(value: str) -> str:
    match = re.search(r"(?:SEC\s+)?[0-9*][0-9*\s-]{5,}[0-9*]", value, re.IGNORECASE)
    return match.group(0).strip() if match else value.strip()


def extract_bank_pdf_transactions(data: bytes) -> str:
    """Extract canonical transaction CSV from bank statement PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError as exc:
        raise ValueError("PDF流水解析需要安装 pdfplumber") from exc

    rows: list[dict[str, Any]] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        all_tables: list[list[list[str]]] = []
        owner_name = "涉案账户所有人"
        owner_account = ""
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            for line in page_text.splitlines()[:6]:
                m_owner = re.search(r"(?:户名|客户名称|姓名|账户名)[:：\s]+([\u4e00-\u9fff\w]+)", line)
                if m_owner and owner_name == "涉案账户所有人":
                    owner_name = m_owner.group(1).strip()
                m_acc = re.search(r"(?:账号|卡号|账户|银行账号)[:：\s]+(\d{10,25})", line)
                if m_acc and not owner_account:
                    owner_account = m_acc.group(1).strip()

            tables = page.extract_tables()
            for t in tables:
                if t and len(t) >= 2:
                    all_tables.append(t)

    if not all_tables:
        raise ValueError("PDF未检测到流水表格；若为拍照扫描版，请确认文字方向清晰并使用图片导入")

    for table in all_tables:
        header_index = None
        for i, row in enumerate(table):
            cleaned = [_clean_cell(c) for c in row if c]
            if any(h in cleaned for h in ("交易时间", "交易日期", "日期", "记账日", "记账日期")) and any(
                h in cleaned for h in ("金额", "交易金额", "发生额", "贷方发生额", "借方发生额", "收/支", "收支")
            ):
                header_index = i
                break
        if header_index is None:
            continue

        headers = [_clean_cell(c) for c in table[header_index]]
        mapping = _column_mapping(headers)

        credit_col = next((idx for idx, h in enumerate(headers) if any(n in h for n in ("贷方发生额", "贷方金额", "收入金额", "贷方"))), None)
        debit_col = next((idx for idx, h in enumerate(headers) if any(n in h for n in ("借方发生额", "借方金额", "支出金额", "借方"))), None)

        for raw_row in table[header_index + 1 :]:
            if not raw_row or len(raw_row) < len(headers):
                continue
            cells = [_clean_cell(c) for c in raw_row]

            dt_idx = mapping.get("date_time")
            if dt_idx is None or dt_idx >= len(cells):
                continue
            parsed_time = _parse_datetime(cells[dt_idx])
            if parsed_time is None:
                continue

            direction = None
            amount = None
            if credit_col is not None and debit_col is not None:
                c_val = _number(cells[credit_col]) if credit_col < len(cells) else None
                d_val = _number(cells[debit_col]) if debit_col < len(cells) else None
                if c_val and c_val > 0:
                    direction = "收入"
                    amount = c_val
                elif d_val and d_val > 0:
                    direction = "支出"
                    amount = d_val
            elif "amount" in mapping and mapping["amount"] < len(cells):
                amount = _number(cells[mapping["amount"]])
                if "direction" in mapping and mapping["direction"] < len(cells):
                    direction = cells[mapping["direction"]]
                else:
                    direction = "收入"

            if amount is None or amount <= 0:
                continue

            incoming = direction in {"收入", "入", "贷"}
            counterparty = cells[mapping["counterparty"]] if "counterparty" in mapping and mapping["counterparty"] < len(cells) else ""
            if not counterparty:
                counterparty = "交易对手"

            payee = owner_name if incoming else counterparty
            payer = counterparty if incoming else owner_name
            serial = cells[mapping["serial"]] if "serial" in mapping and mapping["serial"] < len(cells) and cells[mapping["serial"]] else f"TX-PDF-{len(rows)+1}"

            rows.append({
                "transaction_id": serial,
                "date": parsed_time.strftime("%Y-%m-%d"),
                "time": parsed_time.strftime("%H:%M:%S"),
                "payer": payer,
                "payer_account": "" if incoming else owner_account,
                "payee": payee,
                "payee_account": owner_account if incoming else "",
                "amount": f"{amount:.2f}",
                "remark": cells[mapping["remark"]] if "remark" in mapping and mapping["remark"] < len(cells) else "PDF流水导入",
            })

    if not rows:
        raise ValueError("未能从 PDF 表格中解析出有效银行流水记录")

    output = io.StringIO(newline="")
    writer = csv.DictWriter(output, fieldnames=["transaction_id", "date", "time", "payer", "payer_account", "payee", "payee_account", "amount", "remark"])
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue()
